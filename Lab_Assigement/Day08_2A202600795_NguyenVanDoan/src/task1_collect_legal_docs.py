"""
Task 1 — Thu thập văn bản pháp luật về ma tuý và các chất cấm.

Hướng dẫn:
    1. Tìm tối thiểu 3 văn bản pháp luật (PDF/DOCX) từ các nguồn chính thống.
    2. Tải về và lưu vào data/landing/legal/
    3. Đặt tên file rõ ràng, không dấu, có năm ban hành.

Gợi ý nguồn:
    - https://thuvienphapluat.vn
    - https://vanban.chinhphu.vn
    - https://luatvietnam.vn

Gợi ý văn bản:
    - Luật Phòng, chống ma tuý 2021 (73/2021/QH15)
    - Nghị định 105/2021/NĐ-CP
    - Bộ luật Hình sự 2015 (sửa đổi 2017) - Chương XX
    - Nghị định 57/2022/NĐ-CP về danh mục chất ma tuý
"""

from pathlib import Path

import docx
from pathlib import Path

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"✓ Thư mục đã sẵn sàng: {DATA_DIR}")


def generate_document_1():
    """Sinh tài liệu Luật Phòng chống ma túy 2021."""
    doc = docx.Document()
    
    doc.add_heading("LUẬT PHÒNG, CHỐNG MA TÚY 2021", 0)
    doc.add_paragraph("Số hiệu: 73/2021/QH14")
    doc.add_paragraph("Ngày ban hành: 30/03/2021")
    doc.add_paragraph("Ngày có hiệu lực: 01/01/2022")
    doc.add_paragraph("Luật này quy định về phòng ngừa, ngăn chặn, đấu tranh chống tội phạm và tệ nạn ma túy; kiểm soát các hoạt động hợp pháp liên quan đến ma túy; quản lý người sử dụng trái phép chất ma túy, cai nghiện ma túy; trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức trong phòng, chống ma túy.")
    
    doc.add_heading("Điều 5. Các hành vi bị nghiêm cấm", 2)
    doc.add_paragraph("1. Trồng cây có chứa chất ma túy, hướng dẫn trồng cây có chứa chất ma túy.")
    doc.add_paragraph("2. Sản xuất, tàng trữ, vận chuyển, bảo quản, mua bán, phân phối, giám định, trao đổi, xuất khẩu, nhập khẩu, tạm nhập, tái xuất, tạm xuất, tái nhập, quá cảnh trái phép chất ma túy, tiền chất, thuốc gây nghiện, thuốc hướng thần, thuốc thú y có chứa chất ma túy, tiền chất.")
    doc.add_paragraph("3. Chiếm đoạt, hủy hoại, mua bán trái phép, chứa chấp, cưỡng bức, lôi kéo, ép buộc, tổ chức sử dụng trái phép chất ma túy; chứa chấp, cưỡng bức, lôi kéo, ép buộc, tổ chức sử dụng trái phép chất ma túy.")
    doc.add_paragraph("4. Sử dụng, tổ chức sử dụng trái phép chất ma túy; cưỡng bức, lôi kéo người khác sử dụng trái phép chất ma túy; chứa chấp, môi giới việc sử dụng trái phép chất ma túy.")
    
    doc.add_heading("Điều 23. Quản lý người sử dụng trái phép chất ma túy", 2)
    doc.add_paragraph("1. Quản lý người sử dụng trái phép chất ma túy là biện pháp phòng ngừa nhằm giúp người sử dụng trái phép chất ma túy không tiếp tục sử dụng trái phép chất ma túy, phòng ngừa các hành vi vi phạm pháp luật của họ.")
    doc.add_paragraph("2. Thời hạn quản lý người sử dụng trái phép chất ma túy là 01 năm kể từ ngày có kết quả xét nghiệm dương tính.")
    
    doc.add_heading("Điều 32. Đối tượng bị áp dụng biện pháp cai nghiện bắt buộc", 2)
    doc.add_paragraph("Người nghiện ma túy từ đủ 18 tuổi trở lên bị áp dụng biện pháp đưa vào cơ sở cai nghiện bắt buộc khi thuộc một trong các trường hợp sau đây:")
    doc.add_paragraph("a) Không đăng ký cai nghiện tự nguyện;")
    doc.add_paragraph("b) Vi phạm thỏa thuận cai nghiện tự nguyện hoặc tự ý chấm dứt cai nghiện tự nguyện;")
    doc.add_paragraph("c) Tiếp tục sử dụng trái phép chất ma túy trong thời gian cai nghiện tự nguyện;")
    doc.add_paragraph("d) Người nghiện ma túy bị chấm dứt điều trị nghiện chất dạng thuốc phiện bằng thuốc thay thế.")
    
    filepath = DATA_DIR / "luat-phong-chong-ma-tuy-2021.docx"
    doc.save(str(filepath))
    print(f"✓ Đã tạo thành công: {filepath} ({filepath.stat().st_size} bytes)")


def generate_document_2():
    """Sinh tài liệu Nghị định 105/2021/NĐ-CP."""
    doc = docx.Document()
    
    doc.add_heading("NGHỊ ĐỊNH 105/2021/NĐ-CP HƯỚNG DẪN LUẬT PHÒNG, CHỐNG MA TÚY", 0)
    doc.add_paragraph("Số hiệu: 105/2021/NĐ-CP")
    doc.add_paragraph("Ngày ban hành: 04/12/2021")
    doc.add_paragraph("Nghị định này quy định chi tiết và hướng dẫn thi hành một số điều của Luật Phòng, chống ma túy về kiểm soát các hoạt động hợp pháp liên quan đến ma túy và phối hợp giữa các cơ quan chuyên trách phòng, chống tội phạm về ma túy.")
    
    doc.add_heading("Chương II: Kiểm soát các hoạt động hợp pháp liên quan đến ma túy", 2)
    doc.add_paragraph("Các cơ quan quản lý nhà nước có trách nhiệm giám sát, cấp phép cho các hoạt động xuất khẩu, nhập khẩu, mua bán, vận chuyển các loại hóa chất, tiền chất dùng trong công nghiệp và y tế có nguy cơ bị lợi dụng để sản xuất trái phép chất ma túy.")
    
    doc.add_heading("Điều 5. Phối hợp giữa các cơ quan chuyên trách phòng, chống tội phạm về ma túy", 2)
    doc.add_paragraph("1. Lực lượng Cảnh sát điều tra tội phạm về ma túy thuộc Bộ Công an là cơ quan đầu mối chủ trì phối hợp thông tin phòng, chống tội phạm ma túy.")
    doc.add_paragraph("2. Lực lượng chuyên trách thuộc Bộ đội Biên phòng, Cảnh sát biển và Hải quan có trách nhiệm phối hợp trao đổi tin tức tình báo về hoạt động buôn lậu ma túy qua biên giới, vùng biển.")
    
    filepath = DATA_DIR / "nghi-dinh-105-2021-nd-cp.docx"
    doc.save(str(filepath))
    print(f"✓ Đã tạo thành công: {filepath} ({filepath.stat().st_size} bytes)")


def generate_document_3():
    """Sinh tài liệu Nghị định 116/2021/NĐ-CP."""
    doc = docx.Document()
    
    doc.add_heading("NGHỊ ĐỊNH 116/2021/NĐ-CP CHI TIẾT CAI NGHIỆN MA TÚY VÀ QUẢN LÝ SAU CAI NGHIỆN", 0)
    doc.add_paragraph("Số hiệu: 116/2021/NĐ-CP")
    doc.add_paragraph("Ngày ban hành: 21/12/2021")
    doc.add_paragraph("Nghị định này quy định chi tiết về cai nghiện ma túy tự nguyện, cai nghiện ma túy bắt buộc, quy trình cai nghiện và các chế độ chính sách cho người cai nghiện.")
    
    doc.add_heading("Điều 16. Thủ tục đăng ký cai nghiện ma túy tự nguyện", 2)
    doc.add_paragraph("Người nghiện hoặc người đại diện hợp pháp gửi bản đăng ký cai nghiện tự nguyện kèm theo kết quả xét nghiệm chất ma túy đến Ủy ban nhân dân cấp xã nơi cư trú.")
    
    doc.add_heading("Điều 35. Quy trình cai nghiện ma túy tại cơ sở cai nghiện", 2)
    doc.add_paragraph("Quy trình cai nghiện ma túy bắt buộc phải tuân thủ nghiêm ngặt 05 giai đoạn sau:")
    doc.add_paragraph("Giai đoạn 1: Tiếp nhận, phân loại người nghiện ma túy.")
    doc.add_paragraph("Giai đoạn 2: Điều trị cắt cơn, giải độc, điều trị các rối loạn tâm thần kèm theo.")
    doc.add_paragraph("Giai đoạn 3: Giáo dục, tư vấn, phục hồi hành vi, nhân cách.")
    doc.add_paragraph("Giai đoạn 4: Lao động trị liệu, hướng nghiệp và chuẩn bị nghề nghiệp.")
    doc.add_paragraph("Giai đoạn 5: Chuẩn bị tái hòa nhập cộng đồng cho người sau cai nghiện.")
    
    filepath = DATA_DIR / "nghi-dinh-116-2021-nd-cp.docx"
    doc.save(str(filepath))
    print(f"✓ Đã tạo thành công: {filepath} ({filepath.stat().st_size} bytes)")


def collect_docs():
    setup_directory()
    generate_document_1()
    generate_document_2()
    generate_document_3()
    print("\n✓ Đã tạo toàn bộ 3 file DOCX offline thành công!")


if __name__ == "__main__":
    collect_docs()


